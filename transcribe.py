##  Simple program to take an audio file and transcribe it to unformatted text
##  File for transcription will be expected in ~/transcription/audio
##  Output transcription will be stored in ~/transcription/output
##  Default output will be .docx, with .txt and .odt as alternatives
##
##  required libraries:
##      python-docx        (pip3 install python-docx)
##      SpeechRecognition  (pip3 install SpeechRecognition)
##      sys                (default library)
##      os                 (default library)
##
##
##  adjust folderPath to suit your needs

from docx import Document
import speech_recognition as sr
import sys
import os

##The filename will ultimately be passed by the web page on upload. For now,
## it will be passed from the command line.
document = Document()
folderPath = "/home/canthus13/transcription/"
filename = sys.argv[1]
recordIt = sr.Recognizer()
fileToTranscribe = sr.AudioFile(folderPath + "audio/" + filename)

## verify that the file type is .wav, .AIFF/.AIFF-C, or FLAC

if filename.lower().endswith(('.wav', '.aiff', 'flac')):
    pass
else:
    ##convert incompatible file type to .wav... TODO
    pass

with fileToTranscribe as source:
    audio = recordIt.record(source)

output = recordIt.recognize_google(audio)
print(output)

## Create output file.

document.add_heading(filename, 0)
p = document.add_paragraph(output)
document.save(folderPath + "output/" + filename + ".docx")
